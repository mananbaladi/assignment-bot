import io
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

# Logo path — same folder as this script's parent
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO_PATH = os.path.join(BASE_DIR, "szabist_logo.png")


# ─── WORD ────────────────────────────────────────────────────────────────────

def export_word(info, content):
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── HEADER ───────────────────────────────────────────────────────────────
    hdr = section.header
    # Clear default blank paragraph
    for p in hdr.paragraphs:
        p._element.getparent().remove(p._element)

    # Row 1: Logo (left) + University name (right of logo)
    hdr_p1 = hdr.add_paragraph()
    hdr_p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add logo image inline
    if os.path.exists(LOGO_PATH):
        run_img = hdr_p1.add_run()
        run_img.add_picture(LOGO_PATH, height=Cm(1.2))
    # University name next to logo
    run_name = hdr_p1.add_run("  Shaheed Zulfikar Ali Bhutto Institute of Science &Technology")
    run_name.font.size = Pt(11)
    run_name.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
    run_name.bold = False

    # Row 2: Grey bar — "AI AND ROBOTICS" full width centered
    hdr_p2 = hdr.add_paragraph()
    hdr_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Grey shading on paragraph
    pPr = hdr_p2._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'BFBFBF')
    pPr.append(shd)
    # Border on paragraph
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '1')
        b.set(qn('w:color'), '000000')
        pBdr.append(b)
    pPr.append(pBdr)
    dept_run = hdr_p2.add_run(info['department'])
    dept_run.bold = True
    dept_run.font.size = Pt(11)
    dept_run.font.color.rgb = RGBColor(0, 0, 0)

    # ── BODY ─────────────────────────────────────────────────────────────────

    # Blank spacer
    doc.add_paragraph()

    # Total Marks / Obtained Marks — right aligned, bold, with underline value
    def marks_row(label, value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rl = p.add_run(f"{label}:   ")
        rl.bold = True
        rl.font.size = Pt(11)
        rv = p.add_run(value)
        rv.bold = True
        rv.underline = True
        rv.font.size = Pt(11)
        return p

    marks_row("Total Marks", f"  {info['total_marks']}  ")
    marks_row("Obtained Marks", "  ______  ")

    doc.add_paragraph()
    doc.add_paragraph()

    # Subject — large, bold, centered
    subj_p = doc.add_paragraph()
    subj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subj_p.add_run(info['subject'])
    sr.bold = True
    sr.font.size = Pt(26)

    # Assignment # — bold, underline, centered
    asn_p = doc.add_paragraph()
    asn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = asn_p.add_run(f"Assignment # {info['assignment_no']}")
    ar.bold = True
    ar.underline = True
    ar.font.size = Pt(16)

    doc.add_paragraph()

    # Last date of Submission — bold, underline, centered, two lines
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Last date of Submission:\n{info['submission_date']}")
    dr.bold = True
    dr.underline = True
    dr.font.size = Pt(10)

    # Spacer lines
    for _ in range(4):
        doc.add_paragraph()

    # Info rows: label bold, value normal, full-width underline via bottom border
    def info_row(label, value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rl = p.add_run(f"{label}:  ")
        rl.bold = True
        rl.font.size = Pt(11)
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        # Bottom border = underline across full width
        pPr2 = p._p.get_or_add_pPr()
        pBdr2 = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '000000')
        pBdr2.append(bot)
        pPr2.append(pBdr2)
        doc.add_paragraph()  # spacer after each row

    info_row("Submitted To", info['submitted_to'])
    info_row("Student Name", info['student_name'])
    info_row("Reg. Number",  info['reg_number'])

    # ── CONTENT ──────────────────────────────────────────────────────────────
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        clean = line.lstrip('#').lstrip('-').strip()
        if not clean:
            doc.add_paragraph()
            continue

        if clean.isupper() and len(clean.replace(' ', '')) > 3:
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
        else:
            p = doc.add_paragraph(clean)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                r.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)

    # ── FOOTER ───────────────────────────────────────────────────────────────
    ftr = section.footer
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Top border on footer
    fpPr = fp._p.get_or_add_pPr()
    fpBdr = OxmlElement('w:pBdr')
    ftop = OxmlElement('w:top')
    ftop.set(qn('w:val'), 'single')
    ftop.set(qn('w:sz'), '6')
    ftop.set(qn('w:space'), '1')
    ftop.set(qn('w:color'), '000000')
    fpBdr.append(ftop)
    fpPr.append(fpBdr)

    dept_code = info.get('dept_code', 'DR&AI')
    # Three-column footer using tabs
    fp.add_run(dept_code)
    # Tab stops for center and right
    pPr3 = fp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_center = OxmlElement('w:tab')
    tab_center.set(qn('w:val'), 'center')
    tab_center.set(qn('w:pos'), '4680')  # center of page
    tabs.append(tab_center)
    tab_right = OxmlElement('w:tab')
    tab_right.set(qn('w:val'), 'right')
    tab_right.set(qn('w:pos'), '9360')
    tabs.append(tab_right)
    pPr3.append(tabs)

    fp.add_run(f"\t{info['batch']}\tSZABIST-ISB").font.size = Pt(9)
    for r in fp.runs:
        r.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── PDF ─────────────────────────────────────────────────────────────────────

class AssignmentPDF(FPDF):
    def __init__(self, info):
        super().__init__()
        self.info = info

    def header(self):
        # Row 1: Logo + University name
        if os.path.exists(LOGO_PATH):
            self.image(LOGO_PATH, x=10, y=8, h=12)
        self.set_xy(26, 8)
        self.set_font('Helvetica', '', 12)
        self.set_text_color(0, 51, 153)
        self.cell(0, 12, 'Shaheed Zulfikar Ali Bhutto Institute of Science &Technology',
                  new_x='LMARGIN', new_y='NEXT')

        # Row 2: Grey bar full width
        self.set_fill_color(191, 191, 191)
        self.set_draw_color(0, 0, 0)
        self.set_text_color(0, 0, 0)
        self.set_font('Helvetica', 'B', 11)
        self.cell(0, 8, self.info['department'], border=1, align='C',
                  fill=True, new_x='LMARGIN', new_y='NEXT')
        self.ln(3)

    def footer(self):
        self.set_y(-15)
        # Top line
        self.set_draw_color(0, 0, 0)
        self.set_line_width(0.4)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(1)
        self.set_font('Helvetica', '', 9)
        self.set_text_color(0, 0, 0)
        dept_code = self.info.get('dept_code', 'DR&AI')
        # Three columns
        w = (self.w - 20) / 3
        self.cell(w, 6, dept_code, align='L')
        self.cell(w, 6, self.info['batch'], align='C')
        self.cell(w, 6, 'SZABIST-ISB', align='R')


def export_pdf(info, content):
    pdf = AssignmentPDF(info)
    pdf.set_margins(10, 32, 10)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    pdf.set_text_color(0, 0, 0)

    # Total Marks / Obtained Marks — right, bold, underlined value
    pdf.set_font('Helvetica', 'B', 11)
    pdf.cell(0, 7, f"Total Marks:    {info['total_marks']}",
             align='R', new_x='LMARGIN', new_y='NEXT')
    pdf.set_x(pdf.l_margin)
    pdf.cell(0, 7, "Obtained Marks:    ______",
             align='R', new_x='LMARGIN', new_y='NEXT')
    pdf.set_x(pdf.l_margin)
    pdf.ln(8)

    # Subject — large bold centered
    pdf.set_font('Helvetica', 'B', 26)
    pdf.cell(0, 14, info['subject'], align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.set_x(pdf.l_margin)

    # Assignment # — bold underline centered
    pdf.set_font('Helvetica', 'BU', 16)
    pdf.cell(0, 10, f"Assignment # {info['assignment_no']}", align='C',
             new_x='LMARGIN', new_y='NEXT')
    pdf.set_x(pdf.l_margin)
    pdf.ln(4)

    # Last date — bold underline centered
    pdf.set_font('Helvetica', 'BU', 10)
    pdf.cell(0, 6, "Last date of Submission:", align='C',
             new_x='LMARGIN', new_y='NEXT')
    pdf.set_x(pdf.l_margin)
    pdf.cell(0, 6, info['submission_date'], align='C',
             new_x='LMARGIN', new_y='NEXT')
    pdf.set_x(pdf.l_margin)
    pdf.ln(16)

    # Info rows with full-width underline
    def info_row(label, value):
        pdf.set_font('Helvetica', 'B', 11)
        pdf.cell(38, 7, f"{label}:")
        pdf.set_font('Helvetica', '', 11)
        pdf.cell(0, 7, f"  {value}", new_x='LMARGIN', new_y='NEXT')
        # Draw underline across full width
        y = pdf.get_y()
        pdf.set_draw_color(0, 0, 0)
        pdf.set_line_width(0.3)
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.set_x(pdf.l_margin)
        pdf.ln(4)

    info_row("Submitted To", info['submitted_to'])
    info_row("Student Name", info['student_name'])
    info_row("Reg. Number",  info['reg_number'])
    pdf.ln(6)

    # Body
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
        clean = line.lstrip('#').lstrip('-').strip()
        if not clean:
            pdf.ln(3)
            continue

        if clean.isupper() and len(clean.replace(' ', '')) > 3:
            pdf.set_font('Helvetica', 'B', 11)
            pdf.ln(3)
            pdf.set_x(pdf.l_margin)
            pdf.cell(0, 7, clean, new_x='LMARGIN', new_y='NEXT')
        else:
            pdf.set_font('Helvetica', '', 11)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 5.5, clean, align='J')
            pdf.set_x(pdf.l_margin)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.read()
